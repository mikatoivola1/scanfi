"""
Generate Word document for ScanFi Software Audit Process.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_audit_document():
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('ScanFi Software Audit Process', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Pre-Seed Investment Due Diligence & Regulatory Compliance')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Document info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Document Version: 1.0\n')
    info.add_run(f'Date: {datetime.now().strftime("%Y-%m-%d")}\n')
    info.add_run('Classification: Confidential')

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Code Quality & Development Standards',
        '3. Security Assessment',
        '4. GDPR & EU Data Protection Compliance',
        '5. Google Play Store Requirements',
        '6. Apple App Store Requirements',
        '7. Audit Checklists',
        '8. Remediation Tracking',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_heading('1.1 Audit Scope', level=2)
    doc.add_paragraph(
        'This audit process evaluates ScanFi across four dimensions:'
    )
    scope_items = [
        'Technical Quality: Code architecture, maintainability, scalability',
        'Security Posture: Vulnerability assessment, data protection',
        'Regulatory Compliance: GDPR, ePrivacy Directive, EU requirements',
        'Platform Compliance: Google Play and Apple App Store policies',
    ]
    for item in scope_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('1.2 Current Architecture', level=2)

    # Architecture table
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = ['Component', 'Technology', 'Risk Level']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    arch_data = [
        ('Backend', 'FastAPI (Python)', 'Low'),
        ('Frontend', 'Vanilla JS PWA', 'Low'),
        ('Data Sources', 'Open Food Facts, Edamam, MyMemory', 'Medium'),
        ('Hosting', 'Render.com', 'Low'),
        ('Authentication', 'None (anonymous)', 'Low'),
        ('Data Storage', 'None (stateless)', 'Low'),
    ]
    for i, (comp, tech, risk) in enumerate(arch_data, 1):
        table.rows[i].cells[0].text = comp
        table.rows[i].cells[1].text = tech
        table.rows[i].cells[2].text = risk

    doc.add_paragraph()

    doc.add_heading('1.3 Key Findings', level=2)
    doc.add_paragraph(
        'ScanFi presents a favorable risk profile for investment due to its privacy-by-design '
        'architecture. No personal data is stored on servers, minimizing GDPR obligations and '
        'data breach risks.'
    )

    # Strengths
    doc.add_paragraph('Strengths:', style='Intense Quote')
    strengths = [
        'No user data stored (GDPR data minimization by design)',
        'No authentication system (no credential breach risk)',
        'Stateless architecture (scalable, simple)',
        'Modern tech stack (FastAPI, PWA)',
    ]
    for s in strengths:
        doc.add_paragraph(s, style='List Bullet')

    # Areas for improvement
    doc.add_paragraph('Areas Requiring Attention:', style='Intense Quote')
    improvements = [
        'Privacy policy not yet published (blocker for app stores)',
        'Terms of service not yet published',
        'No automated testing infrastructure',
        'No CI/CD pipeline',
        'Basic logging only',
    ]
    for s in improvements:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_page_break()

    # Section 2: Code Quality
    doc.add_heading('2. Code Quality & Development Standards', level=1)

    doc.add_heading('2.1 Version Control Requirements', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Requirement', 'Status', 'Evidence Required']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    vc_data = [
        ('Git repository with full history', 'To verify', 'GitHub/GitLab URL'),
        ('Protected main branch', 'To verify', 'Branch protection rules'),
        ('Signed commits', 'To verify', 'GPG key verification'),
        ('Meaningful commit messages', 'To verify', 'Commit history review'),
    ]
    for i, (req, status, evidence) in enumerate(vc_data, 1):
        table.rows[i].cells[0].text = req
        table.rows[i].cells[1].text = status
        table.rows[i].cells[2].text = evidence

    doc.add_paragraph()

    doc.add_heading('2.2 Testing Requirements', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Test Type', 'Coverage Target', 'Current State']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    test_data = [
        ('Unit tests', '80%', 'Not implemented'),
        ('Integration tests', 'Critical paths', 'Not implemented'),
        ('E2E tests', 'Happy paths', 'Not implemented'),
        ('Security tests', 'OWASP Top 10', 'Not implemented'),
    ]
    for i, (test, target, state) in enumerate(test_data, 1):
        table.rows[i].cells[0].text = test
        table.rows[i].cells[1].text = target
        table.rows[i].cells[2].text = state

    doc.add_paragraph()

    doc.add_heading('2.3 CI/CD Pipeline Requirements', level=2)
    doc.add_paragraph('Minimum pipeline stages required:')
    stages = [
        'Lint: Python (flake8, black) and JavaScript (eslint)',
        'Security: Dependency scan (pip-audit), SAST (bandit), secret detection',
        'Test: Unit tests with minimum 80% coverage gate',
        'Build: Docker image build and vulnerability scan',
        'Deploy: Staging deployment, smoke tests, production (manual approval)',
    ]
    for s in stages:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_page_break()

    # Section 3: Security Assessment
    doc.add_heading('3. Security Assessment', level=1)

    doc.add_heading('3.1 OWASP Top 10 Checklist', level=2)

    table = doc.add_table(rows=11, cols=4)
    table.style = 'Table Grid'
    headers = ['#', 'Vulnerability', 'Risk', 'Status']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    owasp_data = [
        ('A01', 'Broken Access Control', 'Low', 'N/A - No auth'),
        ('A02', 'Cryptographic Failures', 'Low', 'HTTPS enforced'),
        ('A03', 'Injection', 'Medium', 'Needs validation'),
        ('A04', 'Insecure Design', 'Medium', 'Threat model needed'),
        ('A05', 'Security Misconfiguration', 'Medium', 'Headers needed'),
        ('A06', 'Vulnerable Components', 'High', 'Audit needed'),
        ('A07', 'Authentication Failures', 'N/A', 'No authentication'),
        ('A08', 'Software Integrity Failures', 'Medium', 'SRI needed'),
        ('A09', 'Logging & Monitoring', 'High', 'Not implemented'),
        ('A10', 'SSRF', 'Low', 'Review needed'),
    ]
    for i, (num, vuln, risk, status) in enumerate(owasp_data, 1):
        table.rows[i].cells[0].text = num
        table.rows[i].cells[1].text = vuln
        table.rows[i].cells[2].text = risk
        table.rows[i].cells[3].text = status

    doc.add_paragraph()

    doc.add_heading('3.2 Security Headers Required', level=2)
    headers_list = [
        'X-Content-Type-Options: nosniff',
        'X-Frame-Options: DENY',
        'X-XSS-Protection: 1; mode=block',
        'Referrer-Policy: strict-origin-when-cross-origin',
        'Content-Security-Policy: (configured for app resources)',
        'Permissions-Policy: camera=(self), geolocation=()',
    ]
    for h in headers_list:
        doc.add_paragraph(h, style='List Bullet')

    doc.add_heading('3.3 Penetration Testing', level=2)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    headers = ['Test Type', 'Frequency', 'Provider', 'Status']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    pentest_data = [
        ('Automated DAST', 'Weekly', 'OWASP ZAP', 'Not configured'),
        ('Manual pentest', 'Pre-launch', 'Third-party', 'Not scheduled'),
        ('Bug bounty', 'Post-launch', 'Platform TBD', 'Not established'),
    ]
    for i, (test, freq, provider, status) in enumerate(pentest_data, 1):
        table.rows[i].cells[0].text = test
        table.rows[i].cells[1].text = freq
        table.rows[i].cells[2].text = provider
        table.rows[i].cells[3].text = status

    doc.add_page_break()

    # Section 4: GDPR Compliance
    doc.add_heading('4. GDPR & EU Data Protection Compliance', level=1)

    doc.add_heading('4.1 Data Processing Assessment', level=2)

    doc.add_paragraph('Data Categories Processed by ScanFi:', style='Intense Quote')

    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    headers = ['Data Category', 'Collected', 'Stored', 'Shared', 'Legal Basis']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    data_categories = [
        ('Barcode scans', 'Yes', 'No', 'Yes (APIs)', 'Legitimate interest'),
        ('Camera feed', 'Yes (local)', 'No', 'No', 'Consent'),
        ('Device language', 'Yes', 'Local only', 'No', 'Contract'),
        ('IP address', 'Server logs', 'Temporary', 'No', 'Legitimate interest'),
        ('Product data', 'Retrieved', 'No', 'No', 'N/A (public)'),
    ]
    for i, row_data in enumerate(data_categories, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('4.2 Compliance Status', level=2)
    doc.add_paragraph(
        'ScanFi demonstrates strong GDPR compliance through data minimization by design. '
        'No personal data is stored on servers, eliminating most data protection obligations.'
    )

    doc.add_heading('4.3 Required Legal Documents', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Document', 'Status', 'Location']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    legal_docs = [
        ('Privacy Policy', 'REQUIRED', '/privacy route'),
        ('Terms of Service', 'REQUIRED', '/terms route'),
        ('Cookie Policy', 'Required if cookies used', 'Part of Privacy Policy'),
        ('Third-Party DPA List', 'REQUIRED', 'Privacy Policy'),
    ]
    for i, (doc_name, status, location) in enumerate(legal_docs, 1):
        table.rows[i].cells[0].text = doc_name
        table.rows[i].cells[1].text = status
        table.rows[i].cells[2].text = location

    doc.add_paragraph()

    doc.add_heading('4.4 User Rights (GDPR Article 12-23)', level=2)

    rights = [
        ('Right to access', 'N/A - No personal data stored'),
        ('Right to rectification', 'N/A - No personal data stored'),
        ('Right to erasure', 'N/A - No personal data stored'),
        ('Right to data portability', 'N/A - No personal data stored'),
        ('Right to object', 'Process documented in privacy policy'),
    ]
    for right, status in rights:
        doc.add_paragraph(f'{right}: {status}', style='List Bullet')

    doc.add_page_break()

    # Section 5: Google Play
    doc.add_heading('5. Google Play Store Requirements', level=1)

    doc.add_heading('5.1 Developer Account', level=2)
    account_reqs = [
        'Google Play Developer account ($25 one-time fee)',
        'D-U-N-S number (required for organization accounts)',
        'Valid payment method',
    ]
    for req in account_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('5.2 Data Safety Form', level=2)
    doc.add_paragraph('Required declarations for ScanFi:')

    safety_items = [
        'Data Types Collected: Barcode data (not personal)',
        'Camera Usage: Used locally only, video never leaves device',
        'Data Sharing: Barcode shared with third-party APIs for product lookup',
        'Data Security: Encrypted in transit (HTTPS)',
        'Data Deletion: N/A - no data stored',
    ]
    for item in safety_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.3 Store Listing Requirements', level=2)

    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    headers = ['Asset', 'Specification', 'Status']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    listing_data = [
        ('App icon', '512x512 PNG', 'To create'),
        ('Feature graphic', '1024x500 PNG', 'To create'),
        ('Screenshots', 'Min 2, various sizes', 'To create'),
        ('Short description', '80 chars max', 'To write'),
        ('Full description', '4000 chars max', 'To write'),
        ('Privacy policy URL', 'Required', 'To create'),
        ('Category', 'Food & Drink', 'To select'),
        ('Content rating', 'IARC questionnaire', 'To complete'),
    ]
    for i, (asset, spec, status) in enumerate(listing_data, 1):
        table.rows[i].cells[0].text = asset
        table.rows[i].cells[1].text = spec
        table.rows[i].cells[2].text = status

    doc.add_page_break()

    # Section 6: Apple App Store
    doc.add_heading('6. Apple App Store Requirements', level=1)

    doc.add_heading('6.1 Developer Account', level=2)
    apple_reqs = [
        'Apple Developer Program ($99/year)',
        'D-U-N-S number (required for organization)',
        'Legal entity verification',
    ]
    for req in apple_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('6.2 App Privacy Details', level=2)
    doc.add_paragraph('Privacy "nutrition label" for App Store:')
    privacy_items = [
        'Data Linked to You: None',
        'Data Not Linked to You: None',
        'Data Used to Track You: None',
        'Camera Usage Description: "ScanFi needs camera access to scan product barcodes"',
    ]
    for item in privacy_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6.3 PWA Considerations', level=2)
    doc.add_paragraph(
        'As a Progressive Web App, ScanFi has limited functionality on iOS Safari. '
        'Consider creating a native wrapper using Capacitor for full App Store presence.'
    )

    doc.add_page_break()

    # Section 7: Checklists
    doc.add_heading('7. Audit Checklists', level=1)

    doc.add_heading('7.1 Pre-Seed Investor Due Diligence', level=2)

    doc.add_paragraph('Code Quality:', style='Intense Quote')
    code_items = [
        'Clean git history with meaningful commits',
        'Consistent code style (linting passes)',
        'No hardcoded secrets in repository',
        'Dependencies up-to-date and secure',
        'README with setup instructions',
        'Architecture documentation',
    ]
    for item in code_items:
        doc.add_paragraph(f'☐ {item}')

    doc.add_paragraph('Security:', style='Intense Quote')
    security_items = [
        'HTTPS enforced everywhere',
        'Security headers implemented',
        'No injection vulnerabilities',
        'Dependency vulnerability scan clean',
        'API rate limiting implemented',
    ]
    for item in security_items:
        doc.add_paragraph(f'☐ {item}')

    doc.add_paragraph('Compliance:', style='Intense Quote')
    compliance_items = [
        'Privacy policy published',
        'Terms of service published',
        'GDPR compliance documented',
        'Third-party data processing documented',
    ]
    for item in compliance_items:
        doc.add_paragraph(f'☐ {item}')

    doc.add_page_break()

    # Section 8: Remediation
    doc.add_heading('8. Remediation Tracking', level=1)

    doc.add_heading('8.1 Priority Matrix', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Priority', 'Criteria', 'Timeline']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    priority_data = [
        ('P0 - Critical', 'Security vulnerability, legal blocker', 'Immediate'),
        ('P1 - High', 'Required for store submission', '1-2 weeks'),
        ('P2 - Medium', 'Investor concern, best practice', '1 month'),
        ('P3 - Low', 'Nice to have, future improvement', 'Backlog'),
    ]
    for i, (priority, criteria, timeline) in enumerate(priority_data, 1):
        table.rows[i].cells[0].text = priority
        table.rows[i].cells[1].text = criteria
        table.rows[i].cells[2].text = timeline

    doc.add_paragraph()

    doc.add_heading('8.2 Current Action Items', level=2)

    table = doc.add_table(rows=11, cols=5)
    table.style = 'Table Grid'
    headers = ['ID', 'Issue', 'Priority', 'Owner', 'Status']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    action_items = [
        ('001', 'Implement security headers', 'P1', '', 'Open'),
        ('002', 'Create privacy policy', 'P1', '', 'Open'),
        ('003', 'Create terms of service', 'P1', '', 'Open'),
        ('004', 'Add structured logging', 'P2', '', 'Open'),
        ('005', 'Implement API rate limiting', 'P2', '', 'Open'),
        ('006', 'Add unit tests (80% coverage)', 'P2', '', 'Open'),
        ('007', 'Set up CI/CD pipeline', 'P2', '', 'Open'),
        ('008', 'Dependency vulnerability scan', 'P1', '', 'Open'),
        ('009', 'Create store listing assets', 'P1', '', 'Open'),
        ('010', 'Complete data safety forms', 'P1', '', 'Open'),
    ]
    for i, (id_, issue, priority, owner, status) in enumerate(action_items, 1):
        table.rows[i].cells[0].text = id_
        table.rows[i].cells[1].text = issue
        table.rows[i].cells[2].text = priority
        table.rows[i].cells[3].text = owner
        table.rows[i].cells[4].text = status

    doc.add_page_break()

    # Appendix
    doc.add_heading('Appendix A: Third-Party Services', level=1)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Service', 'Purpose', 'Data Shared']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    services = [
        ('Open Food Facts', 'Product data', 'Barcode'),
        ('MyMemory', 'Translation', 'Text strings'),
        ('Edamam', 'Nutrition data', 'Barcode, product name'),
        ('Render.com', 'Hosting', 'Server logs'),
    ]
    for i, (service, purpose, data) in enumerate(services, 1):
        table.rows[i].cells[0].text = service
        table.rows[i].cells[1].text = purpose
        table.rows[i].cells[2].text = data

    doc.add_paragraph()

    doc.add_heading('Appendix B: Cost Estimates', level=1)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Item', 'One-time', 'Recurring']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    costs = [
        ('Google Play Developer', '$25', '-'),
        ('Apple Developer Program', '-', '$99/year'),
        ('Legal review (privacy docs)', '$500-2,000', '-'),
        ('Penetration test (optional)', '$2,000-5,000', 'Annual'),
        ('Monitoring tools', '-', '$0-50/month'),
    ]
    for i, (item, onetime, recurring) in enumerate(costs, 1):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = onetime
        table.rows[i].cells[2].text = recurring

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Total Pre-Launch Estimate: ').bold = True
    p.add_run('$525-2,125 (one-time) + $99-149/year')

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('This document should be reviewed and updated quarterly or upon significant changes.')
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    output_path = 'C:/scanfi/docs/ScanFi_Software_Audit_Process.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_audit_document()
